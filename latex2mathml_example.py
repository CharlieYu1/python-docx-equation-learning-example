from docx import Document
from lxml import etree
import latex2mathml.converter


def latex_to_word(latex_input):
    mathml = latex2mathml.converter.convert(latex_input)
    tree = etree.fromstring(mathml)
    xslt = etree.parse(
        'MML2OMML.XSL'
        )
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


document = Document()

p = document.add_paragraph()
word_math = latex_to_word(r"\sum_{i=1}^{10}{\frac{\sigma_{zp,i}}{E_i} kN")
p._element.append(word_math)

p = document.add_paragraph()
p.add_run('before formula ')
p._element.append(word_math)
p.add_run(' after formula ')

p = document.add_paragraph()
word_math = latex_to_word(r"\sum_{n=1}^{10}{n^{2}}")
p._element.append(word_math)

p = document.add_paragraph()
word_math = latex_to_word(r"\sqrt{\frac{a}{b}}")
p._element.append(word_math)


p = document.add_paragraph()
word_math = latex_to_word(r"\begin{matrix*}[r]a & b \\ c & d \end{matrix*}")
p._element.append(word_math)

document.save('demo.docx')